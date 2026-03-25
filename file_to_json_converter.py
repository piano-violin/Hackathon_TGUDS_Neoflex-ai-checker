"""
Универсальный конвертер файлов в JSON формат для тестовых заданий

Этот модуль преобразует ответы пользователей из различных форматов (TXT, PDF, DOCX и др.)
в единый структурированный JSON-формат, подходящий для последующей обработки и анализа.

Основные возможности:
- Поддержка различных форматов входных файлов
- Автоматическое распознавание структуры тестовых заданий
- Приведение ответов к единому формату независимо от исходного формата файла
- Обработка ошибок и валидация данных

Поддерживаемые форматы:
- Текстовые: TXT, MD, SQL
- Документы: PDF, DOCX, DOC
- Табличные: XLSX

Пример использования:
    from file_to_json_converter import convert_to_json
    json_data = convert_to_json("in/7837378.txt")
    
Структура выходного JSON:
{
    "answers": {
        "1": "Содержимое ответа на задание 1",
        "2": "Содержимое ответа на задание 2",
        ...
    }
}

Дополнительные зависимости:
    Для XLSX: pip install openpyxl
    Для DOCX: pip install python-docx
    Для PDF: установите Poppler (pdftotext)
"""

import json
import os
import subprocess
import re
import shutil
import zlib
from pathlib import Path
from typing import Dict, Any, List, Optional
import logging

logger = logging.getLogger(__name__)

# Флаги наличия сторонних библиотек
HAS_OPENPYXL = False
HAS_PYTHON_DOCX = False

try:
    from openpyxl import load_workbook
    HAS_OPENPYXL = True
except ImportError:
    pass

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    pass


def read_txt(filepath: str) -> Dict[str, Any]:
    """
    Читает содержимое текстовых файлов (TXT, MD, SQL).
    
    Args:
        filepath: Путь к файлу для чтения
    
    Returns:
        Словарь с содержимым файла
    """
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Обработка SQL-файлов
        if filepath.lower().endswith('.sql'):
            statements = [s.strip() for s in content.split(';') if s.strip()]
            return {'type': 'sql', 'content': content, 'statements': statements}
        
        # Обработка Markdown-файлов
        if filepath.lower().endswith('.md'):
            lines = content.split('\n')
            headers = [l for l in lines if l.startswith('#')]
            return {'type': 'markdown', 'content': content, 'headers': headers}
        
        # Обычный текст
        return {'type': 'text', 'content': content}
    
    except Exception as e:
        return {'type': 'text', 'error': str(e), 'content': ''}


def read_pdf(filepath: str) -> Dict[str, Any]:
    """
    Конвертирует PDF-файл в текст с использованием pdftotext.
    
    Args:
        filepath: Путь к PDF-файлу
    
    Returns:
        Словарь с текстовым содержимым и метаданными
    """
    try:
        metadata = {}

        # Основной путь: Poppler (если установлен в системе)
        if shutil.which('pdftotext'):
            result = subprocess.run(
                ['pdftotext', '-layout', filepath, '-'],
                capture_output=True,
                timeout=60
            )
            content = _decode_content(result.stdout)

            if shutil.which('pdfinfo'):
                meta_result = subprocess.run(
                    ['pdfinfo', filepath],
                    capture_output=True,
                    timeout=30
                )
                meta_output = _decode_content(meta_result.stdout)
                for line in meta_output.split('\n'):
                    if ':' in line:
                        key, value = line.split(':', 1)
                        metadata[key.strip()] = value.strip()
        else:
            # Fallback без внешних CLI: извлекаем текст из PDF-потоков и ToUnicode CMap.
            content = _extract_pdf_text_python(filepath)

        return {
            'type': 'pdf',
            'content': content,
            'metadata': metadata,
            'pages': metadata.get('Pages', 'unknown')
        }

    except Exception as e:
        return {
            'type': 'pdf',
            'error': str(e),
            'content': '',
            'metadata': {},
            'pages': 'unknown'
        }


def _extract_pdf_text_python(filepath: str) -> str:
    """
    Резервный извлекатель текста из простых PDF без внешних утилит.
    Работает с FlateDecode потоками и ToUnicode CMap.
    """
    with open(filepath, 'rb') as f:
        pdf_bytes = f.read()

    objects = _extract_pdf_objects(pdf_bytes)
    streams = []
    for obj_num, raw_obj in objects:
        stream_data = _extract_and_decompress_stream(raw_obj)
        if stream_data is not None:
            streams.append((obj_num, raw_obj, stream_data))

    cmap = {}
    for _, _, data in streams:
        text = _decode_content(data)
        if 'beginbfchar' in text:
            cmap.update(_parse_tounicode_bfchar(text))

    extracted_lines: list[str] = []
    for _, _, data in streams:
        text = _decode_content(data)
        if 'BT' not in text:
            continue
        extracted_lines.extend(_extract_text_lines_from_content_stream(text, cmap))

    return "\n".join(line for line in extracted_lines if line.strip()).strip()


def _extract_pdf_objects(pdf_bytes: bytes) -> list[tuple[int, bytes]]:
    pattern = re.compile(rb'(\d+)\s+0\s+obj(.*?)endobj', re.S)
    objects: list[tuple[int, bytes]] = []
    for match in pattern.finditer(pdf_bytes):
        obj_num = int(match.group(1))
        raw_obj = match.group(2)
        objects.append((obj_num, raw_obj))
    return objects


def _extract_and_decompress_stream(raw_obj: bytes) -> bytes | None:
    match = re.search(rb'stream\r?\n(.*?)\r?\nendstream', raw_obj, re.S)
    if not match:
        return None

    stream_bytes = match.group(1)
    if b'/FlateDecode' in raw_obj:
        try:
            return zlib.decompress(stream_bytes)
        except Exception:
            return None
    return stream_bytes


def _parse_tounicode_bfchar(cmap_text: str) -> Dict[str, str]:
    mapping: Dict[str, str] = {}
    for src_hex, dst_hex in re.findall(r'<([0-9A-Fa-f]+)>\s*<([0-9A-Fa-f]+)>', cmap_text):
        try:
            src_bytes = bytes.fromhex(src_hex)
            dst_bytes = bytes.fromhex(dst_hex)
            if not src_bytes or not dst_bytes:
                continue
            mapping[src_bytes.hex().upper()] = dst_bytes.decode('utf-16-be', errors='ignore')
        except Exception:
            continue
    return mapping


def _extract_text_lines_from_content_stream(stream_text: str, cmap: Dict[str, str]) -> list[str]:
    lines: list[str] = []
    for match in re.finditer(r'\[(.*?)\]\s*TJ|<([0-9A-Fa-f]+)>\s*Tj', stream_text, flags=re.S):
        if match.group(1):
            chunks = re.findall(r'<([0-9A-Fa-f]+)>', match.group(1))
        else:
            chunks = [match.group(2)]

        decoded = ''.join(_decode_pdf_hex_chunk(chunk, cmap) for chunk in chunks).strip()
        if decoded:
            lines.append(decoded)
    return lines


def _decode_pdf_hex_chunk(hex_text: str, cmap: Dict[str, str]) -> str:
    if not hex_text:
        return ''

    hex_text = hex_text.upper()
    out: list[str] = []
    for i in range(0, len(hex_text), 2):
        code = hex_text[i:i + 2]
        if len(code) < 2:
            continue
        out.append(cmap.get(code, ''))
    return ''.join(out)


def _decode_content(content_bytes: bytes) -> str:
    """
    Безопасное декодирование байтов в строку.
    
    Args:
        content_bytes: Байтовая строка для декодирования
    
    Returns:
        Декодированная строка с заменой нечитаемых символов
    """
    try:
        return content_bytes.decode('utf-8', errors='replace')
    except:
        try:
            return content_bytes.decode('cp1251', errors='replace')
        except:
            return content_bytes.decode('latin-1', errors='replace')


def read_docx(filepath: str) -> Dict[str, Any]:
    """
    Читает содержимое DOCX файлов.
    
    Args:
        filepath: Путь к DOCX-файлу
    
    Returns:
        Словарь с текстовым содержимым
    """
    if HAS_PYTHON_DOCX:
        return _read_docx_python(filepath)
    else:
        return _read_docx_pandoc(filepath)


def _read_docx_python(filepath: str) -> Dict[str, Any]:
    """
    Внутренняя реализация чтения DOCX через python-docx.
    
    Args:
        filepath: Путь к DOCX-файлу
    
    Returns:
        Структурированные данные из документа
    """
    try:
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs]
        tables = []
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'type': 'docx',
            'content': '\n'.join(paragraphs),
            'paragraphs': paragraphs,
            'tables': tables,
            'method': 'python-docx'
        }
    except Exception as e:
        return {'type': 'docx', 'error': str(e), 'content': ''}


def _read_docx_pandoc(filepath: str) -> Dict[str, Any]:
    """
    Внутренняя реализация чтения DOCX через pandoc.
    
    Args:
        filepath: Путь к DOCX-файлу
    
    Returns:
        Текстовое содержимое документа
    """
    try:
        result = subprocess.run(
            ['pandoc', filepath, '-t', 'plain'],
            capture_output=True,
            text=True,
            timeout=60
        )
        return {
            'type': 'docx',
            'content': result.stdout,
            'method': 'pandoc'
        }
    except Exception as e:
        return {'type': 'docx', 'error': str(e), 'content': ''}


def read_doc(filepath: str) -> Dict[str, Any]:
    """
    Читает содержимое старых DOC файлов через LibreOffice.
    
    Args:
        filepath: Путь к DOC-файлу
    
    Returns:
        Словарь с текстовым содержимым
    """
    try:
        # Создаем временную директорию
        output_dir = os.path.join(os.path.dirname(filepath), 'doc_convert')
        os.makedirs(output_dir, exist_ok=True)
        
        # Конвертируем DOC в DOCX
        base_name = Path(filepath).stem
        docx_path = os.path.join(output_dir, f'{base_name}.docx')
        
        subprocess.run(
            ['soffice', '--headless', '--convert-to', 'docx', 
             '--outdir', output_dir, filepath],
            capture_output=True,
            timeout=120
        )
        
        # Читаем сконвертированный файл
        if os.path.exists(docx_path):
            result = _read_docx_python(docx_path)
            result['original_format'] = 'doc'
            return result
        
        return {'type': 'doc', 'error': 'Conversion failed', 'content': ''}
    
    except Exception as e:
        return {'type': 'doc', 'error': str(e), 'content': ''}


def read_xlsx(filepath: str) -> Dict[str, Any]:
    """
    Читает содержимое XLSX файлов через openpyxl.
    
    Args:
        filepath: Путь к XLSX-файлу
    
    Returns:
        Словарь с данными таблицы
    """
    if not HAS_OPENPYXL:
        return {
            'type': 'xlsx', 
            'error': 'openpyxl not installed. Run: pip install openpyxl',
            'content': ''
        }
    
    try:
        wb = load_workbook(filepath, data_only=True)
        
        result = {
            'type': 'xlsx',
            'sheets': {},
            'sheet_names': wb.sheetnames
        }
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            # Преобразуем в список списков
            data = []
            for row in ws.iter_rows(values_only=True):
                row_data = [str(cell) if cell is not None else '' for cell in row]
                data.append(row_data)
            
            # Первая строка как заголовки
            headers = data[0] if data else []
            
            result['sheets'][sheet_name] = {
                'data': data,
                'headers': headers,
                'rows': len(data),
                'columns': len(headers)
            }
        
        return result
    
    except Exception as e:
        return {'type': 'xlsx', 'error': str(e), 'content': ''}


def ocr_image(image_path: str) -> str:
    """
    Извлекает текст из изображения с помощью Tesseract OCR.
    Работает как в Windows, так и в Linux.
    
    Args:
        image_path: Путь к изображению
    
    Returns:
        Распознанный текст или сообщение об ошибке
    """
    try:
        from PIL import Image
        import pytesseract
        
        # Убираем жесткий путь - pytesseract найдет tesseract автоматически
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang='rus+eng')
        
        return text.strip()
    
    except ImportError as e:
        logger.error(f"Отсутствует зависимость для OCR: {str(e)}")
        return ""
    except Exception as e:
        logger.error(f"Ошибка OCR при обработке {image_path}: {str(e)}")
        return ""


def neural_parse_test_answers(content: str) -> Dict[str, Dict[str, str]]:
    """
    Нейросетевой парсер для структурирования ответов кандидата на 4 тестовых задания.

    Функция использует модель семейства Qwen2.5-Coder для извлечения ответов из 
    неструктурированного текста (content). Поддерживает обработку пропусков, 
    исправляет типичные ошибки генерации LLM и возвращает результат в формате JSON.

    ### Основная логика:
    1. **Формирование контекста**: Используется `apply_chat_template` для разделения 
       инструкций (System Prompt) и данных пользователя (User Prompt).
    2. **Контроль генерации**: Параметр `repetition_penalty=1.1` предотвращает бесконечное 
       повторение токенов, если модель «зациклится» на одном ответе.
    3. **Точный срез (Slicing)**: Из общего тензора выходных данных вырезается только 
       новая последовательность (ответ), исключая входной промпт.
    4. **Постобработка (Clean-up)**: Реализована защита от «галлюцинаций дублирования». 
       Если модель копирует текст задания №3 в поле задания №4, Python-скрипт 
       автоматически заменяет дубль на пустую строку.

    Args:
        content (str): Текст, полученный после первичного парсинга ответа кандидата.

    Returns:
        Dict[str, Dict[str, str]]: Словарь с ключом 'answers', содержащий 4 строковых поля (1-4).
        Пример: {"answers": {"1": "текст", "2": "", "3": "текст", "4": "текст"}}

    Raises:
        JSONDecodeError: Если модель не смогла сформировать валидный JSON (обрабатывается внутри).
        Exception: При сбоях инференса (логируется, вызывается резервный parse_test_answers).
    """
    # Безопасный fallback: если локальная LLM недоступна, используем надежный rule-based парсер.
    try:
        from model import qwen_model, qwen_tokenizer
        import torch
    except Exception:
        return parse_test_answers(content)

    if not qwen_model or not qwen_tokenizer:
        return parse_test_answers(content)
    
    try:
        # 1. Промпт без лишнего, но с четким указанием на разделение
        prompt = [
            {"role": "system", "content": "Ты — технический парсер. Раздели текст на 4 задания (1, 2, 3, 4). Если ответа на номер нет — пиши пустую строку. Верни только JSON."},
            {"role": "user", "content": f"Внимательно изучи текст кандидата и извлеки ответы: и пояснения\n\n{content}\n\nВерни результат строго в формате JSON: {{\"answers\": {{\"1\": \"...\", \"2\": \"...\", \"3\": \"...\", \"4\": \"...\"}}}}"}
        ]

        text = qwen_tokenizer.apply_chat_template(prompt, tokenize=False, add_generation_prompt=True)
        model_inputs = qwen_tokenizer([text], return_tensors="pt").to(qwen_model.device)

        with torch.no_grad():
            generated_ids = qwen_model.generate(
                **model_inputs,
                max_new_tokens=2048,
                do_sample=False,
                repetition_penalty=1.1, # Защита от зацикливания
                pad_token_id=qwen_tokenizer.eos_token_id
            )
        
        # 2. Чистый срез
        input_length = model_inputs.input_ids.shape[1]
        response_ids = generated_ids[0][input_length:]
        result = qwen_tokenizer.decode(response_ids, skip_special_tokens=True).strip()

        # 3. Парсинг
        json_start = result.find('{')
        json_end = result.rfind('}') + 1
        
        if json_start != -1:
            parsed_result = json.loads(result[json_start:json_end])
            raw_answers = parsed_result.get("answers", {})
            
            # Очистка от дублей
            clean_answers = {}
            for i in range(1, 5):
                curr_key = str(i)
                # Берем текст из JSON
                val = str(raw_answers.get(curr_key, raw_answers.get(i, ""))).strip()
                
                # Сравниваем с предыдущим ответом
                prev_val = clean_answers.get(str(i-1), "")
                
                # Если текущий ответ идентичен предыдущему — зануляем его
                if i > 1 and val == prev_val and val != "":
                    clean_answers[curr_key] = ""
                else:
                    clean_answers[curr_key] = val
                    
            return {"answers": clean_answers}
        
        return parse_test_answers(content)
        
    except Exception as e:
        logger.error(f"Ошибка: {e}")
        return parse_test_answers(content)


def parse_test_answers(content: str) -> Dict[str, Dict[str, str]]:
    """
    Надежный парсер, специально разработанный для тестовых заданий.
    
    Этот парсер:
    1. Ищет только ключевые разделители (1., 2., 3., 4.)
    2. Работает с двумя форматами: простая нумерация и "Тестовое задание № X"
    3. Гарантирует 4 ответа (даже если некоторые пустые)
    4. Сохраняет структуру для последующего анализа нейросетью
    
    Args:
        content: Текстовое содержимое файла
    
    Returns:
        Словарь с ответами, структурированный по номерам заданий (1-4)
    """
    # Шаг 1: сохраняем структуру строк, только мягко чистим шум.
    text = re.sub(r'\r\n|\r', '\n', content or "")
    text = re.sub(r'_{10,}', '', text)
    text = text.strip()

    answers = {str(i): "" for i in range(1, 5)}
    if not text:
        return {"answers": answers}

    markers = _find_task_markers(text)
    if not markers:
        return {"answers": answers}

    for i, marker in enumerate(markers):
        task_num = marker["task_num"]
        start_content = marker["header_end"]
        end_content = markers[i + 1]["start"] if i + 1 < len(markers) else len(text)
        task_text = text[start_content:end_content].strip()
        if 1 <= task_num <= 4:
            answers[str(task_num)] = _clean_task_text(task_text)

    return {"answers": answers}


def _find_task_markers(text: str) -> List[Dict[str, int]]:
    """
    Ищет заголовки заданий в безопасном порядке:
    1) явные "1 задание)" / "Тестовое задание № 1"
    2) fallback на "1)." в начале строки
    """
    markers: List[Dict[str, int]] = []
    seen_positions = set()

    heading_patterns = [
        r'(?im)^\s*(?:тестовое\s+)?задани[ея]\s*№?\s*([1-4])\s*[\)\].:\-]?',
        r'(?im)^\s*([1-4])\s*задани[ея]\s*[\)\].:\-]?',
    ]
    for pattern in heading_patterns:
        for match in re.finditer(pattern, text):
            task_num = int(match.group(1))
            pos = match.start()
            if pos in seen_positions:
                continue
            seen_positions.add(pos)
            markers.append(
                {
                    "task_num": task_num,
                    "start": pos,
                    "header_end": match.end(),
                }
            )

    # Если нашли часть заданий по явным заголовкам, добираем недостающие номера
    # простым паттерном "N) ..." только после последнего явного заголовка.
    if markers:
        found_nums = {m["task_num"] for m in markers}
        last_explicit_start = max(m["start"] for m in markers)
        for missing_num in range(1, 5):
            if missing_num in found_nums:
                continue
            simple_pattern = rf'(?m)^\s*{missing_num}\s*[)\].:\-]\s+'
            for match in re.finditer(simple_pattern, text):
                if match.start() <= last_explicit_start:
                    continue
                pos = match.start()
                if pos in seen_positions:
                    continue
                seen_positions.add(pos)
                markers.append(
                    {
                        "task_num": missing_num,
                        "start": pos,
                        "header_end": match.end(),
                    }
                )
                break
    else:
        for match in re.finditer(r'(?m)^\s*([1-4])\s*[)\].:\-]\s+', text):
            task_num = int(match.group(1))
            pos = match.start()
            if pos in seen_positions:
                continue
            seen_positions.add(pos)
            markers.append(
                {
                    "task_num": task_num,
                    "start": pos,
                    "header_end": match.end(),
                }
            )

    markers.sort(key=lambda x: x["start"])
    return markers


def _clean_task_text(task_text: str) -> str:
    cleaned = re.sub(r'\n{3,}', '\n\n', task_text).strip()
    return cleaned


def normalize_content(content: str) -> str:
    """
    Нормализует текст перед подачей в LLM.
    Решает типичные проблемы с OCR и форматированием.
    """
    if not content or not isinstance(content, str):
        return ""

    # 1. Единый перенос строк
    text = re.sub(r'\r\n|\r', '\n', content)

    # 2. Убираем лишние пробелы и табуляцию
    text = re.sub(r'[ \t]+', ' ', text)

    # 3. Заменяем длинные дефисы/пробелы на обычные
    text = re.sub(r'[_\u2013\u2014]{8,}', '---', text)  # длинные разделители → ---

    # 4. Исправляем частые OCR-ошибки
    text = text.replace('lndex', 'index') \
               .replace('cl1ent', 'client') \
               .replace('nanager', 'manager') \
               .replace('reg_date –дата', 'reg_date date')

    # 5. Фиксируем сломанные ключевые слова SQL
    text = re.sub(r'SUM\( SALARY\)', 'SUM(SALARY)', text)
    text = re.sub(r'MANAGER_ID= NULL', 'MANAGER_ID IS NULL', text)
    text = re.sub(r'WHERE CITY=\'SEOUL\'', 'WHERE CITY = \'SEOUL\'', text)

    # 6. Удаляем бинарный мусор (если есть)
    try:
        text = text.encode('utf-8', errors='ignore').decode('utf-8')
    except:
        pass

    # 7. Сжимаем множественные пустые строки
    text = re.sub(r'\n\s*\n', '\n', text).strip()

    return text


def convert_to_json(filepath: str) -> Dict[str, Any]:
    """Универсальная функция конвертации"""
    ext = Path(filepath).suffix.lower()
    
    # Получаем содержимое файла в зависимости от формата
    if ext in ['.txt', '.md', '.sql']:
        content = read_txt(filepath).get('content', '')
    elif ext == '.pdf':
        content = read_pdf(filepath).get('content', '')
    elif ext == '.docx':
        content = read_docx(filepath).get('content', '')
    elif ext == '.doc':
        content = read_doc(filepath).get('content', '')
    elif ext == '.xlsx':
        xlsx_data = read_xlsx(filepath)
        if 'sheets' in xlsx_data and xlsx_data['sheet_names']:
            first_sheet = xlsx_data['sheet_names'][0]
            content = "\n".join([" ".join(row) for row in xlsx_data['sheets'][first_sheet]['data']])
        else:
            content = ''
    elif ext in ['.png', '.jpg', '.jpeg']:
        content = ocr_image(filepath)
    else:
        content = ''
    
    # Если не удалось получить содержимое, возвращаем ошибку
    if not content.strip():
        return {"error": "Empty content after extraction", "format": ext}
    
    # Нормализация контента
    content = normalize_content(content)
    
    # Базовый и устойчивый парсер (без зависимости от LLM).
    parsed = parse_test_answers(content)
    answers = parsed.get("answers", {})
    if any(str(v).strip() for v in answers.values()):
        return parsed

    # Если rule-based ничего не извлек, пробуем нейросеть.
    return neural_parse_test_answers(content)


def convert_directory(input_dir: str, output_file: Optional[str] = None) -> List[Dict[str, Any]]:
    """
    Конвертирует все поддерживаемые файлы в указанной директории.
    
    Args:
        input_dir: Путь к директории с входными файлами
        output_file: Опциональный путь для сохранения результатов в JSON файл
    
    Returns:
        Список результатов конвертации для каждого файла
    """
    results = []
    supported_ext = ['.txt', '.pdf', '.docx', '.doc', '.md', '.sql', '.xlsx', '.png', '.jpg', '.jpeg']
    
    # Проверяем существование директории
    if not os.path.isdir(input_dir):
        raise ValueError(f"Directory not found: {input_dir}")
    
    # Обрабатываем файлы
    for filename in os.listdir(input_dir):
        filepath = os.path.join(input_dir, filename)
        
        if os.path.isfile(filepath):
            ext = Path(filename).suffix.lower()
            if ext in supported_ext:
                logger.info(f"Обработка файла: {filename}")
                result = convert_to_json(filepath)
                results.append({
                    "filename": filename,
                    "result": result
                })
    
    # Сохраняем результаты, если указан файл
    if output_file and results:
        os.makedirs(os.path.dirname(os.path.abspath(output_file)), exist_ok=True)
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        logger.info(f"Результаты сохранены в: {output_file}")
    
    return results
